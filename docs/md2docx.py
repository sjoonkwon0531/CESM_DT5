#!/usr/bin/env python3
"""Markdown to DOCX converter using python-docx."""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def parse_inline(paragraph, text):
    """Parse inline markdown (bold, italic, code, math) and add runs."""
    # Pattern: **bold**, *italic*, `code`, $math$
    pattern = re.compile(r'(\*\*(.+?)\*\*)|(\*(.+?)\*)|(`(.+?)`)|(\$(.+?)\$)')
    pos = 0
    for m in pattern.finditer(text):
        # Add text before match
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.size = Pt(11)
        if m.group(2):  # bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.font.size = Pt(11)
        elif m.group(4):  # italic
            run = paragraph.add_run(m.group(4))
            run.italic = True
            run.font.size = Pt(11)
        elif m.group(6):  # code
            run = paragraph.add_run(m.group(6))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        elif m.group(8):  # math
            run = paragraph.add_run(m.group(8))
            run.italic = True
            run.font.size = Pt(11)
        pos = m.end()
    # Remaining text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(11)


def md_to_docx(md_path, docx_path):
    lines = Path(md_path).read_text(encoding='utf-8').split('\n')
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # Horizontal rule
        if stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a thin line via border
            run = p.add_run('─' * 60)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue
        
        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            h = doc.add_heading(level=min(level, 4))
            parse_inline(h, text)
            i += 1
            continue
        
        # Code block
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Add shading via XML
            pPr = p._p.get_or_add_pPr()
            shd = pPr.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): 'F5F5F5'
            })
            pPr.append(shd)
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            continue
        
        # Display math ($$...$$)
        if stripped.startswith('$$'):
            math_lines = [stripped.replace('$$', '')]
            i += 1
            while i < len(lines) and '$$' not in lines[i]:
                math_lines.append(lines[i].strip())
                i += 1
            if i < len(lines):
                math_lines.append(lines[i].strip().replace('$$', ''))
                i += 1
            math_text = ' '.join(l for l in math_lines if l)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(math_text)
            run.italic = True
            run.font.size = Pt(11)
            continue
        
        # Table
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            # Parse table
            rows_data = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                # Skip separator rows
                if cells and all(re.match(r'^[-:]+$', c) for c in cells):
                    continue
                rows_data.append(cells)
            if rows_data:
                ncols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=ncols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row in enumerate(rows_data):
                    for ci, cell_text in enumerate(row):
                        if ci < ncols:
                            cell = table.cell(ri, ci)
                            cell.text = ''
                            p = cell.paragraphs[0]
                            parse_inline(p, cell_text)
                            p.paragraph_format.space_before = Pt(2)
                            p.paragraph_format.space_after = Pt(2)
                            # Bold header row
                            if ri == 0:
                                for run in p.runs:
                                    run.bold = True
                                    run.font.size = Pt(10)
                            else:
                                for run in p.runs:
                                    run.font.size = Pt(10)
                # Header row shading
                for ci in range(ncols):
                    cell = table.cell(0, ci)
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = tcPr.makeelement(qn('w:shd'), {
                        qn('w:val'): 'clear',
                        qn('w:color'): 'auto',
                        qn('w:fill'): 'D9E2F3'
                    })
                    tcPr.append(shd)
            continue
        
        # Blockquote
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            pPr = p._p.get_or_add_pPr()
            shd = pPr.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): 'FFF3CD'
            })
            pPr.append(shd)
            parse_inline(p, text)
            i += 1
            continue
        
        # Bullet list
        if re.match(r'^[-*]\s+', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, text)
            i += 1
            continue
        
        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            parse_inline(p, text)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline(p, stripped)
        i += 1
    
    doc.save(docx_path)
    print(f"✅ {docx_path}")


if __name__ == '__main__':
    base = Path(__file__).parent
    files = [
        ('technical-document-ko.md', 'technical-document-ko.docx'),
        ('technical-document-en.md', 'technical-document-en.docx'),
        ('advisory-board-final-review.md', 'advisory-board-final-review.docx'),
    ]
    for md, docx in files:
        md_to_docx(base / md, base / docx)
